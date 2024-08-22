from flask import Flask, request, jsonify
from flask_sqlalchemy import SQLAlchemy
from sentence_transformers import SentenceTransformer
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
import uuid
from datetime import datetime, timezone
import chromadb
from chromadb.config import Settings
import logging

app = Flask(__name__)
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///documents.db'
db = SQLAlchemy(app)

# Setting up logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

# Initializing ChromaDB client
client = chromadb.Client(Settings())
try:
    collection = client.create_collection(name="my_collection")
    logger.debug("ChromaDB collection 'my_collection' created or accessed successfully.")
except Exception as e:
    logger.error(f"Failed to create or access ChromaDB collection: {e}")

# Initializing the model
model = SentenceTransformer('all-MiniLM-L6-v2')

class Document(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    asset_id = db.Column(db.String(36), unique=True, nullable=False)
    document_name = db.Column(db.Text, nullable=False)
    file_type = db.Column(db.Text, nullable=False)
    timestamp = db.Column(db.DateTime, default=datetime.now(timezone.utc))

    def __repr__(self):
        return f'<Document {self.document_name}>'

def init_db():
    #Initializing the database.
    with app.app_context():
        db.drop_all()  # Drops all tables
        db.create_all()  # Creates all tables
        logger.debug("Database initialized.")

def extract_text(file):
    logger.debug(f"Extracting text from file: {file.filename}")
    if file.filename.endswith('.txt'):
        return file.read().decode('utf-8')
    elif file.filename.endswith('.pdf'):
        pdf_reader = PdfReader(file)
        text = ''
        for page in pdf_reader.pages:
            text += page.extract_text()
        return text
    elif file.filename.endswith('.docx'):
        doc = DocxDocument(file)
        return '\n'.join([para.text for para in doc.paragraphs])
    else:
        logger.warning(f"Unsupported file type: {file.filename}")
        return None

def generate_embeddings(text):
    logger.debug("Generating embeddings.")
    return model.encode(text).tolist()  # Converting numpy array into list for JSON serialization

def generate_asset_id():
    return str(uuid.uuid4())

def prepare_metadata(file_name, file_type):
    return {
        "file_name": file_name,
        "file_type": file_type,
        "timestamp": datetime.now(timezone.utc).isoformat()
    }

def store_embeddings_with_id(embeddings, metadata):
    asset_id = generate_asset_id()
    logger.debug(f"Storing embeddings with asset_id: {asset_id}")
    try:
        collection.add(
            ids=[asset_id],
            embeddings=[embeddings],
            metadatas=[metadata]
        )
        logger.debug("Embeddings and metadata stored in ChromaDB.")
    except Exception as e:
        logger.error(f"Failed to store embeddings in ChromaDB: {e}")

    # Storing metadata in SQLAlchemy
    document = Document(asset_id=asset_id, document_name=metadata['file_name'], file_type=metadata['file_type'])
    db.session.add(document)
    db.session.commit()
    logger.debug(f"Metadata stored in SQLAlchemy with asset_id: {asset_id}")
    
    return asset_id

def update_document(asset_id, new_embeddings, new_metadata):
    logger.debug(f"Updating document with asset_id: {asset_id}")
    try:
        collection.update(
            ids=[asset_id],
            embeddings=[new_embeddings],
            metadatas=[new_metadata]
        )
        logger.debug("Embeddings and metadata updated in ChromaDB.")
    except Exception as e:
        logger.error(f"Failed to update embeddings in ChromaDB: {e}")

    # Updating the document in SQLAlchemy
    document = Document.query.filter_by(asset_id=asset_id).first()
    if document:
        document.document_name = new_metadata['file_name']
        document.file_type = new_metadata['file_type']
        document.timestamp = datetime.now(timezone.utc)
        db.session.commit()
        logger.debug(f"Document updated in SQLAlchemy with asset_id: {asset_id}")
    else:
        logger.warning(f"Document with asset_id {asset_id} not found in SQLAlchemy.")

def get_document_by_id(asset_id):
    logger.debug(f"Fetching document by asset_id: {asset_id}")
    # Retrieving the document from SQLAlchemy
    document = Document.query.filter_by(asset_id=asset_id).first()
    
    if not document:
        return {'error': 'Document not found in SQLAlchemy'}, 404

    # Using the document name to fetch the embeddings
    metadata_filter = {"file_name": document.document_name}

    # Query the ChromaDB collection
    try:
        results = collection.query(
            query_texts=[document.document_name],
            where=metadata_filter,
            n_results=1
        )
        logger.debug("ChromaDB query executed.")
    except Exception as e:
        logger.error(f"Failed to query ChromaDB: {e}")
        return {'error': 'Failed to query ChromaDB'}, 500

    if not results['ids']:
        return {'error': 'Document not found in ChromaDB'}, 404

    result_data = {
        'document': {
            'asset_id': document.asset_id,
            'document_name': document.document_name,
            'file_type': document.file_type,
            'id': document.id,
            'timestamp': document.timestamp.isoformat()
        },
        'metadata': results['metadatas'][0][0]
    }

    return result_data

def delete_document(asset_id):
    logger.debug(f"Deleting document with asset_id: {asset_id}")
    try:
        collection.delete(ids=[asset_id])
        logger.debug("Document deleted from ChromaDB.")
    except Exception as e:
        logger.error(f"Failed to delete document from ChromaDB: {e}")

    # Deleting the document in SQLAlchemy
    document = Document.query.filter_by(asset_id=asset_id).first()
    if document:
        db.session.delete(document)
        db.session.commit()
        logger.debug(f"Document deleted from SQLAlchemy with asset_id: {asset_id}")
    else:
        logger.warning(f"Document with asset_id {asset_id} not found in SQLAlchemy.")

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400

    file = request.files['file']

    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400

    text = extract_text(file)
    if not text:
        return jsonify({'error': 'Unsupported file type'}), 400

    embeddings = generate_embeddings(text)
    metadata = prepare_metadata(file.filename, file.content_type)
    asset_id = store_embeddings_with_id(embeddings, metadata)

    return jsonify({'message': 'File processed', 'asset_id': asset_id}), 201

@app.route('/update/<asset_id>', methods=['PUT'])
def update_file(asset_id):
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400

    file = request.files['file']
    text = extract_text(file)
    if not text:
        return jsonify({'error': 'Unsupported file type'}), 400

    embeddings = generate_embeddings(text)
    metadata = prepare_metadata(file.filename, file.content_type)

    # Updating the document
    update_document(asset_id, embeddings, metadata)
    
    return jsonify({'message': 'Document updated successfully'}), 200

@app.route('/document/<asset_id>', methods=['GET'])
def get_document(asset_id):
    result = get_document_by_id(asset_id)
    return jsonify(result), 200

@app.route('/document/<asset_id>', methods=['DELETE'])
def delete_document_route(asset_id):
    delete_document(asset_id)
    return jsonify({'message': 'Document deleted successfully'}), 200

if __name__ == '__main__':
    init_db()  # Initializing the database
    app.run(debug=True)
